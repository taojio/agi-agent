"""
perception/text_perception.py - 文本感知子模块

实现任务 T029-T033：用户原始文本接收、文档解析、OCR 提取、
文本清洗归一化、多语言翻译对齐。

所有重型依赖采用 try/except 可选导入，保证无依赖环境可导入实例化。
"""
from __future__ import annotations

import logging
import re
import time
import uuid
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Sequence

from agi_agent.core import BaseModule

logger = logging.getLogger("agi_agent.perception")


# ============================ 数据结构 ============================

@dataclass
class RawTextMessage:
    """T029 用户原始文本报文

    保留用户输入的原始报文，不进行任何修改。
    """

    raw_text: str
    source: str            # 来源标识：dialog / api / file 等
    timestamp: float
    message_id: str


@dataclass
class ParsedDocument:
    """T030 文档解析结果"""

    text: str
    sections: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class OCRItem:
    """T031 OCR 单条结果"""

    text: str
    bbox: List[int]            # [x1, y1, x2, y2]
    confidence: float = 0.0


# ============================ T029 ============================

class UserTextReceiver(BaseModule):
    """T029 用户原始文本接收

    事件触发模块。接收来自对话框/接口的原始文本输入，保留原始报文，
    并向下游模块提供待处理队列。
    """

    name = "user_text_receiver"
    version = "1.0.0"
    description = "接收用户原始文本输入，保留原始报文"

    def __init__(self) -> None:
        super().__init__()
        self._pending: List[RawTextMessage] = []

    def receive(self, raw_text: str, source: str = "dialog") -> RawTextMessage:
        """接收一条原始文本

        Args:
            raw_text: 用户输入的原始文本
            source: 来源标识（dialog/api/file 等）

        Returns:
            RawTextMessage: 保留原始报文的消息对象
        """
        msg = RawTextMessage(
            raw_text=raw_text,
            source=source,
            timestamp=time.time(),
            message_id=uuid.uuid4().hex,
        )
        self._pending.append(msg)
        logger.debug(
            "接收原始文本 source=%s len=%d id=%s",
            source, len(raw_text), msg.message_id,
        )
        return msg

    def get_pending(self) -> List[RawTextMessage]:
        """获取并清空待处理队列

        Returns:
            List[RawTextMessage]: 当前积累的待处理消息列表
        """
        pending = list(self._pending)
        self._pending.clear()
        return pending


# ============================ T030 ============================

class DocumentParser(BaseModule):
    """T030 文档解析

    动态调度模块。支持 PDF / Word / Excel / 网页等格式的解析，
    自动提取纯文本和章节分段。所有解析依赖（PyPDF/python-docx/openpyxl/BeautifulSoup）
    均为可选，缺失时降级为纯文本读取 + 按段落空行切分。
    """

    name = "document_parser"
    version = "1.0.0"
    description = "解析 PDF/Word/Excel/网页等文档，提取文本与章节"

    def __init__(self) -> None:
        super().__init__()
        self._has_pypdf: bool = False
        self._has_docx: bool = False
        self._has_openpyxl: bool = False
        self._has_bs4: bool = False
        try:
            import PyPDF2  # type: ignore  # noqa: F401
            self._has_pypdf = True
        except Exception:
            try:
                import pypdf  # type: ignore  # noqa: F401
                self._has_pypdf = True
            except Exception:
                pass
        try:
            import docx  # type: ignore  # noqa: F401
            self._has_docx = True
        except Exception:
            pass
        try:
            import openpyxl  # type: ignore  # noqa: F401
            self._has_openpyxl = True
        except Exception:
            pass
        try:
            import bs4  # type: ignore  # noqa: F401
            self._has_bs4 = True
        except Exception:
            pass

    def parse(self, file_path: str) -> ParsedDocument:
        """解析文件路径指向的文档

        Args:
            file_path: 文件路径

        Returns:
            ParsedDocument: 解析结果
        """
        path_lower = file_path.lower()
        try:
            if path_lower.endswith(".pdf"):
                return self._parse_pdf(file_path)
            if path_lower.endswith(".docx"):
                return self._parse_docx(file_path)
            if path_lower.endswith((".xlsx", ".xlsm")):
                return self._parse_xlsx(file_path)
            if path_lower.endswith((".html", ".htm")):
                return self._parse_html_file(file_path)
            return self._parse_plain(file_path)
        except Exception as e:
            logger.warning("文档解析失败 path=%s err=%s，降级纯文本读取", file_path, e)
            return self._parse_plain(file_path)

    def parse_bytes(self, data: bytes, fmt: str) -> ParsedDocument:
        """解析字节流

        Args:
            data: 文档字节流
            fmt: 格式标识 pdf/docx/xlsx/html/txt

        Returns:
            ParsedDocument: 解析结果
        """
        import io
        fmt = (fmt or "").lower()
        try:
            if fmt == "pdf":
                return self._parse_pdf_bytes(io.BytesIO(data))
            if fmt == "docx":
                return self._parse_docx_bytes(io.BytesIO(data))
            if fmt in ("xlsx", "xlsm"):
                return self._parse_xlsx_bytes(io.BytesIO(data))
            if fmt in ("html", "htm"):
                return self._parse_html_bytes(data)
            text = data.decode("utf-8", errors="ignore")
            return ParsedDocument(
                text=text,
                sections=self._split_sections(text),
                metadata={"format": fmt},
            )
        except Exception as e:
            logger.warning("字节流解析失败 fmt=%s err=%s", fmt, e)
            text = data.decode("utf-8", errors="ignore")
            return ParsedDocument(
                text=text,
                sections=self._split_sections(text),
                metadata={"format": fmt, "error": str(e)},
            )

    def _parse_pdf(self, path: str) -> ParsedDocument:
        if not self._has_pypdf:
            logger.debug("PyPDF 不可用，PDF 按二进制降级")
            return self._parse_plain(path)
        with open(path, "rb") as f:
            return self._parse_pdf_bytes(f)

    def _parse_pdf_bytes(self, stream: Any) -> ParsedDocument:
        text_parts: List[str] = []
        try:
            try:
                from PyPDF2 import PdfReader
            except Exception:
                from pypdf import PdfReader  # type: ignore
            reader = PdfReader(stream)
            for page in reader.pages:
                try:
                    text_parts.append(page.extract_text() or "")
                except Exception:
                    continue
        except Exception as e:
            logger.warning("PDF 解析异常 err=%s", e)
        text = "\n".join(text_parts)
        return ParsedDocument(
            text=text,
            sections=self._split_sections(text),
            metadata={"format": "pdf", "pages": len(text_parts)},
        )

    def _parse_docx(self, path: str) -> ParsedDocument:
        if not self._has_docx:
            return self._parse_plain(path)
        import docx  # type: ignore
        d = docx.Document(path)
        return self._docx_to_doc(d)

    def _parse_docx_bytes(self, stream: Any) -> ParsedDocument:
        if not self._has_docx:
            return ParsedDocument(text="", sections=[], metadata={"format": "docx"})
        import docx  # type: ignore
        d = docx.Document(stream)
        return self._docx_to_doc(d)

    def _docx_to_doc(self, d: Any) -> ParsedDocument:
        sections: List[str] = []
        buf: List[str] = []
        for para in d.paragraphs:
            t = para.text or ""
            if not t.strip():
                continue
            style_name = ""
            try:
                style_name = para.style.name if para.style else ""
            except Exception:
                style_name = ""
            if style_name.startswith("Heading"):
                if buf:
                    sections.append("\n".join(buf))
                    buf = []
                sections.append(t)
            else:
                buf.append(t)
        if buf:
            sections.append("\n".join(buf))
        full_text = "\n".join(sections) if sections else "\n".join(p.text for p in d.paragraphs)
        return ParsedDocument(
            text=full_text,
            sections=sections,
            metadata={"format": "docx", "paragraphs": len(d.paragraphs)},
        )

    def _parse_xlsx(self, path: str) -> ParsedDocument:
        if not self._has_openpyxl:
            return self._parse_plain(path)
        import openpyxl  # type: ignore
        wb = openpyxl.load_workbook(path, data_only=True)
        return self._xlsx_to_doc(wb)

    def _parse_xlsx_bytes(self, stream: Any) -> ParsedDocument:
        if not self._has_openpyxl:
            return ParsedDocument(text="", sections=[], metadata={"format": "xlsx"})
        import openpyxl  # type: ignore
        wb = openpyxl.load_workbook(stream, data_only=True)
        return self._xlsx_to_doc(wb)

    def _xlsx_to_doc(self, wb: Any) -> ParsedDocument:
        sections: List[str] = []
        for ws in wb.worksheets:
            rows_text: List[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows_text.append("\t".join(cells))
            if rows_text:
                sections.append(f"[Sheet: {ws.title}]\n" + "\n".join(rows_text))
        full_text = "\n\n".join(sections)
        return ParsedDocument(
            text=full_text,
            sections=sections,
            metadata={"format": "xlsx", "sheets": [ws.title for ws in wb.worksheets]},
        )

    def _parse_html_file(self, path: str) -> ParsedDocument:
        with open(path, "rb") as f:
            return self._parse_html_bytes(f.read())

    def _parse_html_bytes(self, data: bytes) -> ParsedDocument:
        text_raw = data.decode("utf-8", errors="ignore")
        if self._has_bs4:
            from bs4 import BeautifulSoup  # type: ignore
            soup = BeautifulSoup(text_raw, "html.parser")
            for s in soup(["script", "style"]):
                s.decompose()
            text = soup.get_text(separator="\n")
        else:
            text = re.sub(r"<script[^>]*>.*?</script>", "", text_raw, flags=re.S | re.I)
            text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.S | re.I)
            text = re.sub(r"<[^>]+>", "\n", text)
        text = re.sub(r"\n{2,}", "\n\n", text).strip()
        return ParsedDocument(
            text=text,
            sections=self._split_sections(text),
            metadata={"format": "html"},
        )

    def _parse_plain(self, path: str) -> ParsedDocument:
        try:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
        except Exception:
            with open(path, "r", encoding="gbk", errors="ignore") as f:
                text = f.read()
        return ParsedDocument(
            text=text,
            sections=self._split_sections(text),
            metadata={"format": "plain"},
        )

    @staticmethod
    def _split_sections(text: str) -> List[str]:
        """按空行切分段落"""
        if not text:
            return []
        parts = re.split(r"\n\s*\n", text)
        return [p.strip() for p in parts if p.strip()]


# ============================ T031 ============================

class OCRExtractor(BaseModule):
    """T031 OCR 文字提取

    动态调度模块。识别图片中的文字及其坐标。
    paddleocr / pytesseract 均为可选，缺失时返回空结果并 warning。
    """

    name = "ocr_extractor"
    version = "1.0.0"
    description = "识别图片中的文字及坐标"

    def __init__(self) -> None:
        super().__init__()
        self._backend: Optional[str] = None
        self._engine: Any = None
        self._engine_loaded: bool = False
        # 仅探测依赖是否可用，不在构造时初始化引擎（避免触发模型下载）
        try:
            import paddleocr  # type: ignore  # noqa: F401
            self._backend = "paddleocr"
        except Exception:
            pass
        if self._backend is None:
            try:
                import pytesseract  # type: ignore  # noqa: F401
                self._backend = "tesseract"
            except Exception:
                pass
        if self._backend is None:
            logger.warning("未检测到 paddleocr/pytesseract，OCR 将返回空结果")

    def _ensure_engine(self) -> Any:
        """懒加载 OCR 引擎，避免构造时触发网络下载"""
        if self._engine_loaded:
            return self._engine
        self._engine_loaded = True
        try:
            if self._backend == "paddleocr":
                from paddleocr import PaddleOCR  # type: ignore
                self._engine = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
        except Exception as e:
            logger.warning("OCR 引擎初始化失败 err=%s，返回空结果", e)
            self._engine = None
        return self._engine

    def extract(self, image: Any) -> List[OCRItem]:
        """从图像中提取文字

        Args:
            image: numpy 数组 / PIL Image / 文件路径

        Returns:
            List[OCRItem]: OCR 识别结果列表
        """
        if self._backend is None:
            logger.warning("OCR 后端不可用，返回空结果")
            return []
        engine = self._ensure_engine()
        if engine is None:
            return []
        try:
            if self._backend == "paddleocr":
                return self._extract_paddle(engine, image)
            if self._backend == "tesseract":
                return self._extract_tesseract(image)
        except Exception as e:
            logger.warning("OCR 提取失败 err=%s", e)
        return []

    def _extract_paddle(self, engine: Any, image: Any) -> List[OCRItem]:
        import numpy as np
        if isinstance(image, str):
            result = engine.ocr(image, cls=True)
        else:
            if hasattr(image, "convert"):
                image = np.array(image.convert("RGB"))
            result = engine.ocr(image, cls=True)
        items: List[OCRItem] = []
        if not result:
            return items
        for line in result:
            if not line:
                continue
            for item in line:
                try:
                    box = item[0]
                    txt = item[1][0]
                    conf = float(item[1][1])
                    xs = [p[0] for p in box]
                    ys = [p[1] for p in box]
                    bbox = [int(min(xs)), int(min(ys)), int(max(xs)), int(max(ys))]
                    items.append(OCRItem(text=txt, bbox=bbox, confidence=conf))
                except Exception:
                    continue
        return items

    def _extract_tesseract(self, image: Any) -> List[OCRItem]:
        import pytesseract  # type: ignore
        try:
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        except Exception as e:
            logger.warning("tesseract 识别失败 err=%s", e)
            return []
        items: List[OCRItem] = []
        for i, txt in enumerate(data.get("text", [])):
            if not txt or not txt.strip():
                continue
            try:
                left = int(data["left"][i])
                top = int(data["top"][i])
                width = int(data["width"][i])
                height = int(data["height"][i])
                bbox = [left, top, left + width, top + height]
                conf_raw = data["conf"][i]
                conf = float(conf_raw) / 100.0 if conf_raw and int(conf_raw) > 0 else 0.0
                items.append(OCRItem(text=txt, bbox=bbox, confidence=conf))
            except Exception:
                continue
        return items


# ============================ T032 ============================

class TextNormalizer(BaseModule):
    """T032 文本清洗归一化

    动态调度模块。基于正则表达式清洗多余空格、特殊符号、乱码、无效字符，
    统一编码为 UTF-8 规范化文本。
    """

    name = "text_normalizer"
    version = "1.0.0"
    description = "文本清洗归一化（去多余空格、特殊符号、乱码、无效字符）"

    # 控制字符（保留 \t \n）
    _CTRL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
    # 多余空白
    _MULTI_SPACE_RE = re.compile(r"[ \t\u3000]+")
    _MULTI_NEWLINE_RE = re.compile(r"\n{3,}")
    # BOM / 零宽字符
    _BOM_RE = re.compile(r"[\ufeff\u200b\u200c\u200d]")
    # 替换字符 U+FFFD（乱码占位）
    _REPLACEMENT_RE = re.compile(r"\ufffd+")
    # 私有区/不可见符号
    _PRIVATE_USE_RE = re.compile(r"[\ue000-\uf8ff]")

    def __init__(self) -> None:
        super().__init__()

    def normalize(self, text: str) -> str:
        """归一化单条文本

        Args:
            text: 待清洗的原始文本

        Returns:
            str: 清洗后的文本
        """
        if text is None:
            return ""
        if not isinstance(text, str):
            try:
                text = str(text)
            except Exception:
                return ""
        # 统一编码：编解码以丢弃非法字节
        try:
            text = text.encode("utf-8", "ignore").decode("utf-8", "ignore")
        except Exception:
            pass
        # 去 BOM / 零宽字符
        text = self._BOM_RE.sub("", text)
        # 去控制字符
        text = self._CTRL_RE.sub("", text)
        # 去私有区字符
        text = self._PRIVATE_USE_RE.sub("", text)
        # 替换字符归并
        text = self._REPLACEMENT_RE.sub("", text)
        # 全角空格 -> 半角，多空格合并
        text = self._MULTI_SPACE_RE.sub(" ", text)
        # 多换行合并为最多两个
        text = self._MULTI_NEWLINE_RE.sub("\n\n", text)
        # 行首尾空格
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)
        return text.strip()

    def normalize_batch(self, items: Sequence[str]) -> List[str]:
        """批量归一化文本

        Args:
            items: 文本列表

        Returns:
            List[str]: 归一化结果
        """
        return [self.normalize(t) for t in items]


# ============================ T033 ============================

class TranslationAligner(BaseModule):
    """T033 多语言翻译对齐

    动态调度模块。识别外文并将其翻译为中文。
    翻译依赖（translate / requests）为可选，缺失时检测非中文并返回原文 + warning。
    """

    name = "translation_aligner"
    version = "1.0.0"
    description = "多语言识别并翻译为中文"

    # 中日韩字符范围
    _CJK_RE = re.compile(r"[\u3040-\u30ff\u3400-\u4dbf\u4e00-\u9fff\uff00-\uffef]")
    _LATIN_RE = re.compile(r"[A-Za-z]")
    _CYRILLIC_RE = re.compile(r"[\u0400-\u04ff]")
    _KANA_RE = re.compile(r"[\u3040-\u30ff]")
    _HANGUL_RE = re.compile(r"[\uac00-\ud7af]")

    def __init__(self) -> None:
        super().__init__()
        self._has_translate: bool = False
        try:
            import translate  # type: ignore  # noqa: F401
            self._has_translate = True
        except Exception:
            pass
        self._has_requests: bool = False
        try:
            import requests  # type: ignore  # noqa: F401
            self._has_requests = True
        except Exception:
            pass
        if not self._has_translate and not self._has_requests:
            logger.warning("未检测到 translate/requests，翻译将返回原文")

    def detect_language(self, text: str) -> str:
        """检测文本语种

        Args:
            text: 输入文本

        Returns:
            str: 语种代码 zh/en/ja/ko/ru/unknown
        """
        if not text:
            return "unknown"
        sample = text[:500]
        if self._KANA_RE.search(sample):
            return "ja"
        if self._HANGUL_RE.search(sample):
            return "ko"
        if self._CYRILLIC_RE.search(sample):
            return "ru"
        cjk_count = len(self._CJK_RE.findall(sample))
        latin_count = len(self._LATIN_RE.findall(sample))
        if cjk_count > 0 and cjk_count >= latin_count:
            return "zh"
        if latin_count > 0:
            return "en"
        return "unknown"

    def translate(self, text: str, target: str = "zh") -> str:
        """将文本翻译为目标语言（默认中文）

        Args:
            text: 待翻译文本
            target: 目标语种代码

        Returns:
            str: 翻译后文本；无法翻译时返回原文
        """
        if not text:
            return ""
        src = self.detect_language(text)
        if src == target or src == "unknown":
            return text
        if self._has_translate:
            try:
                from translate import Translator  # type: ignore
                translator = Translator(to_lang=target)
                result = translator.translate(text)
                if result and isinstance(result, str):
                    return result
            except Exception as e:
                logger.debug("translate 库翻译失败 err=%s，尝试 requests", e)
        if self._has_requests:
            try:
                import requests  # type: ignore
                resp = requests.post(
                    "https://libretranslate.com/translate",
                    json={
                        "q": text,
                        "source": "auto",
                        "target": target,
                        "format": "text",
                    },
                    timeout=5,
                )
                if resp.status_code == 200:
                    translated = resp.json().get("translatedText")
                    if translated:
                        return str(translated)
            except Exception as e:
                logger.debug("requests 翻译失败 err=%s", e)
        logger.warning("翻译后端不可用，返回原文 src=%s target=%s", src, target)
        return text


__all__ = [
    "RawTextMessage",
    "ParsedDocument",
    "OCRItem",
    "UserTextReceiver",
    "DocumentParser",
    "OCRExtractor",
    "TextNormalizer",
    "TranslationAligner",
]
