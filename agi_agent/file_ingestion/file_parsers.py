import os
import re
from typing import Dict, Optional, Tuple, Any
from datetime import datetime

TEXT_EXTENSIONS = ['.txt', '.md', '.json', '.csv', '.log', '.xml', '.html', '.htm']
PDF_EXTENSIONS = ['.pdf']
DOCX_EXTENSIONS = ['.docx']
AUDIO_EXTENSIONS = ['.mp3', '.wav', '.flac', '.ogg', '.m4a', '.aac']
VIDEO_EXTENSIONS = ['.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv']


class ParsingResult:
    def __init__(self, success: bool, content_type: str, 
                 text_content: str = "", metadata: Dict = None, 
                 error: str = ""):
        self.success = success
        self.content_type = content_type
        self.text_content = text_content
        self.metadata = metadata or {}
        self.error = error


class FileParser:
    def __init__(self):
        self.supported_extensions = []

    def can_parse(self, file_ext: str) -> bool:
        return file_ext.lower() in self.supported_extensions

    def parse(self, file_path: str, binary_content: bytes = None) -> ParsingResult:
        raise NotImplementedError("Subclasses must implement parse method")


class TextParser(FileParser):
    def __init__(self):
        super().__init__()
        self.supported_extensions = TEXT_EXTENSIONS

    def parse(self, file_path: str, binary_content: bytes = None) -> ParsingResult:
        try:
            if binary_content:
                content = binary_content.decode('utf-8', errors='replace')
            else:
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()

            metadata = {
                'file_path': file_path,
                'file_size': len(content.encode('utf-8')),
                'char_count': len(content),
                'line_count': content.count('\n') + 1,
                'parse_time': datetime.now().isoformat()
            }

            return ParsingResult(
                success=True,
                content_type='text',
                text_content=content,
                metadata=metadata
            )

        except Exception as e:
            return ParsingResult(
                success=False,
                content_type='text',
                error=f"Text parsing error: {str(e)}"
            )


class PDFParser(FileParser):
    def __init__(self):
        super().__init__()
        self.supported_extensions = PDF_EXTENSIONS
        self._try_import_pdf()

    def _try_import_pdf(self):
        try:
            import PyPDF2
            self.pdf_reader = PyPDF2.PdfReader
            self._pdf_available = True
        except ImportError:
            self._pdf_available = False

    def parse(self, file_path: str, binary_content: bytes = None) -> ParsingResult:
        if not self._pdf_available:
            return ParsingResult(
                success=False,
                content_type='pdf',
                error="PyPDF2 not installed. Install with: pip install PyPDF2"
            )

        try:
            if binary_content:
                import io
                pdf_file = io.BytesIO(binary_content)
            else:
                pdf_file = open(file_path, 'rb')

            reader = self.pdf_reader(pdf_file)
            content = ""
            for page in reader.pages:
                content += page.extract_text() or ""

            if binary_content:
                pdf_file.close()
            else:
                pdf_file.close()

            metadata = {
                'file_path': file_path,
                'page_count': len(reader.pages),
                'char_count': len(content),
                'parse_time': datetime.now().isoformat()
            }

            return ParsingResult(
                success=True,
                content_type='pdf',
                text_content=content,
                metadata=metadata
            )

        except Exception as e:
            return ParsingResult(
                success=False,
                content_type='pdf',
                error=f"PDF parsing error: {str(e)}"
            )


class DOCXParser(FileParser):
    def __init__(self):
        super().__init__()
        self.supported_extensions = DOCX_EXTENSIONS
        self._try_import_docx()

    def _try_import_docx(self):
        try:
            import docx
            self.docx_module = docx
            self._docx_available = True
        except ImportError:
            self._docx_available = False

    def parse(self, file_path: str, binary_content: bytes = None) -> ParsingResult:
        if not self._docx_available:
            return ParsingResult(
                success=False,
                content_type='docx',
                error="python-docx not installed. Install with: pip install python-docx"
            )

        try:
            if binary_content:
                import io
                docx_file = io.BytesIO(binary_content)
            else:
                docx_file = file_path

            doc = self.docx_module.Document(docx_file)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            metadata = {
                'file_path': file_path,
                'paragraph_count': len(doc.paragraphs),
                'char_count': len(content),
                'parse_time': datetime.now().isoformat()
            }

            return ParsingResult(
                success=True,
                content_type='docx',
                text_content=content,
                metadata=metadata
            )

        except Exception as e:
            return ParsingResult(
                success=False,
                content_type='docx',
                error=f"DOCX parsing error: {str(e)}"
            )


class AudioParser(FileParser):
    def __init__(self):
        super().__init__()
        self.supported_extensions = AUDIO_EXTENSIONS
        self._try_import_audio()

    def _try_import_audio(self):
        try:
            import speech_recognition as sr
            self.sr_module = sr
            self._audio_available = True
        except ImportError:
            self._audio_available = False

    def parse(self, file_path: str, binary_content: bytes = None) -> ParsingResult:
        if not self._audio_available:
            return ParsingResult(
                success=False,
                content_type='audio',
                error="speech_recognition not installed. Install with: pip install SpeechRecognition pydub"
            )

        try:
            recognizer = self.sr_module.Recognizer()
            
            if binary_content:
                import io
                audio_file = io.BytesIO(binary_content)
            else:
                audio_file = file_path

            with self.sr_module.AudioFile(audio_file) as source:
                audio = recognizer.record(source)
            
            try:
                text_content = recognizer.recognize_google(audio, language='zh-CN')
            except self.sr_module.UnknownValueError:
                text_content = "[Audio content could not be transcribed]"
            except self.sr_module.RequestError:
                text_content = "[Speech recognition service unavailable]"

            metadata = {
                'file_path': file_path,
                'content_type': 'audio',
                'duration': source.DURATION if hasattr(source, 'DURATION') else 0,
                'parse_time': datetime.now().isoformat()
            }

            return ParsingResult(
                success=True,
                content_type='audio',
                text_content=text_content,
                metadata=metadata
            )

        except Exception as e:
            return ParsingResult(
                success=False,
                content_type='audio',
                error=f"Audio parsing error: {str(e)}"
            )


class VideoParser(FileParser):
    def __init__(self):
        super().__init__()
        self.supported_extensions = VIDEO_EXTENSIONS

    def parse(self, file_path: str, binary_content: bytes = None) -> ParsingResult:
        try:
            metadata = {
                'file_path': file_path,
                'content_type': 'video',
                'parse_time': datetime.now().isoformat(),
                'processed': False,
                'notes': "Video files require ffmpeg for extraction. Falling back to metadata extraction."
            }

            if os.path.exists(file_path):
                stat = os.stat(file_path)
                metadata['file_size'] = stat.st_size

            text_content = f"[Video file: {os.path.basename(file_path)} - requires ffmpeg for full processing]"

            return ParsingResult(
                success=True,
                content_type='video',
                text_content=text_content,
                metadata=metadata
            )

        except Exception as e:
            return ParsingResult(
                success=False,
                content_type='video',
                error=f"Video parsing error: {str(e)}"
            )


class FileParserManager:
    def __init__(self, logger=None):
        self.logger = logger
        self.parsers = [
            TextParser(),
            PDFParser(),
            DOCXParser(),
            AudioParser(),
            VideoParser()
        ]
        self.parser_cache = {}

    def get_parser(self, file_ext: str) -> Optional[FileParser]:
        file_ext = file_ext.lower()
        
        if file_ext in self.parser_cache:
            return self.parser_cache[file_ext]

        for parser in self.parsers:
            if parser.can_parse(file_ext):
                self.parser_cache[file_ext] = parser
                return parser

        return None

    def parse_file(self, file_path: str, binary_content: bytes = None) -> ParsingResult:
        file_ext = os.path.splitext(file_path)[1].lower()
        parser = self.get_parser(file_ext)

        if not parser:
            return ParsingResult(
                success=False,
                content_type='unknown',
                error=f"No parser available for extension: {file_ext}"
            )

        if self.logger:
            self.logger.info(f"Parsing file: {file_path} with {parser.__class__.__name__}")

        result = parser.parse(file_path, binary_content)
        
        if result.success and self.logger:
            self.logger.info(f"Successfully parsed {file_path}: {len(result.text_content)} chars")
        elif not result.success and self.logger:
            self.logger.error(f"Failed to parse {file_path}: {result.error}")

        return result

    def get_supported_extensions(self) -> Dict[str, str]:
        extensions = {}
        for parser in self.parsers:
            for ext in parser.supported_extensions:
                extensions[ext] = parser.__class__.__name__.replace('Parser', '')
        return extensions